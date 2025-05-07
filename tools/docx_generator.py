#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Generatore di documenti DOCX per il Business Plan Builder

Questo modulo implementa la generazione di documenti DOCX a partire dallo stato
del business plan, creando un documento professionale e ben formattato.
"""

import os
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from state import BusinessPlanState
from config import Config

def generate_docx(state: BusinessPlanState, output_path: str) -> str:
    """Genera un documento DOCX a partire dallo stato del business plan"""
    # Crea un nuovo documento o carica un template
    if os.path.exists(Config.DOCX_TEMPLATE_PATH):
        doc = Document(Config.DOCX_TEMPLATE_PATH)
    else:
        doc = Document()
    
    # Configura gli stili del documento
    setup_document_styles(doc)
    
    # Aggiungi la copertina
    add_cover_page(doc, state)
    
    # Aggiungi il sommario
    doc.add_page_break()
    doc.add_heading("Indice", level=1)
    # Il sommario verrà aggiunto manualmente dall'utente in Word
    doc.add_paragraph("Aggiornare il campo del sommario in Microsoft Word (click destro -> Aggiorna campo)", style="Caption")
    
    # Aggiungi le sezioni del business plan
    for section, subsections in state["outline"].items():
        doc.add_page_break()
        doc.add_heading(section, level=1)
        
        # Recupera i chunk di questa sezione
        section_chunks = get_chunks_for_section(state, section)
        
        # Aggiungi il contenuto della sezione principale
        main_content = get_main_section_content(section_chunks)
        if main_content:
            doc.add_paragraph(main_content)
        
        # Aggiungi le sottosezioni
        for subsection in subsections:
            doc.add_heading(subsection, level=2)
            
            # Recupera i chunk di questa sottosezione
            subsection_chunks = get_chunks_for_subsection(state, section, subsection)
            
            # Aggiungi il contenuto della sottosezione
            subsection_content = get_subsection_content(subsection_chunks)
            if subsection_content:
                doc.add_paragraph(subsection_content)
            else:
                doc.add_paragraph("Contenuto in fase di sviluppo.")
    
    # Aggiungi gli allegati
    if "Appendici" in state["outline"]:
        doc.add_page_break()
        doc.add_heading("Appendici", level=1)
        
        # Aggiungi le appendici
        for appendix in state["outline"]["Appendici"]:
            doc.add_heading(appendix, level=2)
            
            # Recupera i chunk di questa appendice
            appendix_chunks = get_chunks_for_subsection(state, "Appendici", appendix)
            
            # Aggiungi il contenuto dell'appendice
            appendix_content = get_subsection_content(appendix_chunks)
            if appendix_content:
                doc.add_paragraph(appendix_content)
            else:
                doc.add_paragraph("Contenuto in fase di sviluppo.")
    
    # Salva il documento
    doc.save(output_path)
    
    return output_path

def setup_document_styles(doc: Document) -> None:
    """Configura gli stili del documento"""
    # Stile per il titolo
    title_style = doc.styles.add_style("CustomTitle", WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = Config.DEFAULT_FONT
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 0, 90)  # Blu scuro
    
    # Stile per i sottotitoli
    subtitle_style = doc.styles.add_style("CustomSubtitle", WD_STYLE_TYPE.PARAGRAPH)
    subtitle_style.font.name = Config.DEFAULT_FONT
    subtitle_style.font.size = Pt(16)
    subtitle_style.font.italic = True
    subtitle_style.font.color.rgb = RGBColor(70, 70, 70)  # Grigio scuro
    
    # Stile per il testo normale
    normal_style = doc.styles["Normal"]
    normal_style.font.name = Config.DEFAULT_FONT
    normal_style.font.size = Pt(Config.DEFAULT_FONT_SIZE)
    
    # Stile per le intestazioni
    for i in range(1, 4):
        heading_style = doc.styles[f"Heading {i}"]
        heading_style.font.name = Config.DEFAULT_FONT
        heading_style.font.bold = True
        
        if i == 1:
            heading_style.font.size = Pt(18)
            heading_style.font.color.rgb = RGBColor(0, 0, 90)  # Blu scuro
        elif i == 2:
            heading_style.font.size = Pt(16)
            heading_style.font.color.rgb = RGBColor(0, 0, 120)  # Blu medio
        else:
            heading_style.font.size = Pt(14)
            heading_style.font.color.rgb = RGBColor(0, 0, 150)  # Blu chiaro

def add_cover_page(doc: Document, state: BusinessPlanState) -> None:
    """Aggiunge la copertina al documento"""
    # Aggiungi il titolo
    title = doc.add_paragraph(state["document_title"], style="CustomTitle")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Aggiungi il nome dell'azienda
    company = doc.add_paragraph(state["company_name"], style="CustomSubtitle")
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Aggiungi la data di creazione
    date_paragraph = doc.add_paragraph(f"Data: {state['creation_date']}", style="CustomSubtitle")
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Aggiungi la versione
    version_paragraph = doc.add_paragraph(f"Versione: {state['version']}", style="CustomSubtitle")
    version_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def get_chunks_for_section(state: BusinessPlanState, section: str) -> List[Dict[str, Any]]:
    """Recupera i chunk per una sezione specifica"""
    return [chunk for chunk_id, chunk in state["document_chunks"].items() 
            if chunk["section"] == section and chunk["subsection"] is None]

def get_chunks_for_subsection(state: BusinessPlanState, section: str, subsection: str) -> List[Dict[str, Any]]:
    """Recupera i chunk per una sottosezione specifica"""
    return [chunk for chunk_id, chunk in state["document_chunks"].items() 
            if chunk["section"] == section and chunk["subsection"] == subsection]

def get_main_section_content(chunks: List[Dict[str, Any]]) -> str:
    """Estrae il contenuto principale da una lista di chunk"""
    if not chunks:
        return ""
    
    # Ordina i chunk per livello
    chunks.sort(key=lambda x: x["level"])
    
    # Concatena il contenuto
    return "\n\n".join([chunk["content"] for chunk in chunks])

def get_subsection_content(chunks: List[Dict[str, Any]]) -> str:
    """Estrae il contenuto di una sottosezione da una lista di chunk"""
    if not chunks:
        return ""
    
    # Ordina i chunk per livello
    chunks.sort(key=lambda x: x["level"])
    
    # Concatena il contenuto
    return "\n\n".join([chunk["content"] for chunk in chunks])